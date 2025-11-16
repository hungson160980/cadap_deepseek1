from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import matplotlib.pyplot as plt
from io import BytesIO
import base64

class ReportExporter:
    def __init__(self):
        pass
    
    def export_word_report(self, data, include_charts=True):
        """Xuất báo cáo thẩm định dạng Word"""
        doc = Document()
        
        # Tiêu đề
        title = doc.add_heading('BÁO CÁO THẨM ĐỊNH TÍN DỤNG', 0)
        title.alignment = 1
        
        # Thông tin khách hàng
        doc.add_heading('I. THÔNG TIN KHÁCH HÀNG', level=1)
        customer = data.get('customer', {})
        doc.add_paragraph(f"Họ và tên: {customer.get('ho_ten', '')}")
        doc.add_paragraph(f"CCCD/CMND: {customer.get('cccd', '')}")
        doc.add_paragraph(f"Địa chỉ: {customer.get('dia_chi', '')}")
        doc.add_paragraph(f"Số điện thoại: {customer.get('dien_thoai', '')}")
        
        # Thông tin tài chính
        doc.add_heading('II. THÔNG TIN TÀI CHÍNH', level=1)
        financial = data.get('financial', {})
        doc.add_paragraph(f"Tổng nhu cầu vốn: {financial.get('tong_nhu_cau_von', 0):,.0f} VNĐ".replace(",", "."))
        doc.add_paragraph(f"Vốn đối ứng: {financial.get('von_doi_ung', 0):,.0f} VNĐ".replace(",", "."))
        doc.add_paragraph(f"Số tiền vay: {financial.get('so_tien_vay', 0):,.0f} VNĐ".replace(",", "."))
        doc.add_paragraph(f"Lãi suất: {financial.get('lai_suat', 0)}%/năm")
        doc.add_paragraph(f"Thời gian vay: {financial.get('thoi_gian_vay', 0)} tháng")
        doc.add_paragraph(f"Mục đích vay: {financial.get('muc_dich_vay', '')}")
        
        # Chỉ số tài chính
        doc.add_heading('III. CHỈ SỐ TÀI CHÍNH', level=1)
        metrics = data.get('metrics', {})
        doc.add_paragraph(f"Nghĩa vụ trả nợ hàng tháng: {metrics.get('monthly_payment', 0):,.0f} VNĐ".replace(",", "."))
        doc.add_paragraph(f"Tỷ lệ trả nợ (DSR): {metrics.get('dsr_ratio', 0):.1f}%")
        doc.add_paragraph(f"LTV: {metrics.get('ltv', 0):.1f}%")
        doc.add_paragraph(f"Biên an toàn trả nợ: {metrics.get('safety_margin', 0):.1f}%")
        
        # Lưu file vào memory
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.getvalue()
    
    def export_pdf_report(self, data, include_charts=True):
        """Xuất báo cáo thẩm định dạng PDF"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Tiêu đề
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1
        )
        title = Paragraph('BÁO CÁO THẨM ĐỊNH TÍN DỤNG', title_style)
        story.append(title)
        
        # Thông tin khách hàng
        story.append(Paragraph('I. THÔNG TIN KHÁCH HÀNG', styles['Heading2']))
        customer = data.get('customer', {})
        story.append(Paragraph(f"Họ và tên: {customer.get('ho_ten', '')}", styles['Normal']))
        story.append(Paragraph(f"CCCD/CMND: {customer.get('cccd', '')}", styles['Normal']))
        story.append(Paragraph(f"Địa chỉ: {customer.get('dia_chi', '')}", styles['Normal']))
        story.append(Paragraph(f"Số điện thoại: {customer.get('dien_thoai', '')}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Thông tin tài chính
        story.append(Paragraph('II. THÔNG TIN TÀI CHÍNH', styles['Heading2']))
        financial = data.get('financial', {})
        story.append(Paragraph(f"Tổng nhu cầu vốn: {financial.get('tong_nhu_cau_von', 0):,.0f} VNĐ".replace(",", "."), styles['Normal']))
        story.append(Paragraph(f"Vốn đối ứng: {financial.get('von_doi_ung', 0):,.0f} VNĐ".replace(",", "."), styles['Normal']))
        story.append(Paragraph(f"Số tiền vay: {financial.get('so_tien_vay', 0):,.0f} VNĐ".replace(",", "."), styles['Normal']))
        story.append(Paragraph(f"Lãi suất: {financial.get('lai_suat', 0)}%/năm", styles['Normal']))
        story.append(Paragraph(f"Thời gian vay: {financial.get('thoi_gian_vay', 0)} tháng", styles['Normal']))
        story.append(Paragraph(f"Mục đích vay: {financial.get('muc_dich_vay', '')}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Chỉ số tài chính
        story.append(Paragraph('III. CHỈ SỐ TÀI CHÍNH', styles['Heading2']))
        metrics = data.get('metrics', {})
        story.append(Paragraph(f"Nghĩa vụ trả nợ hàng tháng: {metrics.get('monthly_payment', 0):,.0f} VNĐ".replace(",", "."), styles['Normal']))
        story.append(Paragraph(f"Tỷ lệ trả nợ (DSR): {metrics.get('dsr_ratio', 0):.1f}%", styles['Normal']))
        story.append(Paragraph(f"LTV: {metrics.get('ltv', 0):.1f}%", styles['Normal']))
        story.append(Paragraph(f"Biên an toàn trả nợ: {metrics.get('safety_margin', 0):.1f}%", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        
        return buffer.getvalue()