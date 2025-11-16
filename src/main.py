import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
from docx import Document
import re
import json
from io import BytesIO

# =============================================================================
# COMPONENTS & UTILITIES
# =============================================================================

def format_currency(value):
    """Định dạng số tiền với dấu phân cách hàng nghìn"""
    try:
        return f"{float(value):,.0f}".replace(",", ".")
    except (ValueError, TypeError):
        return "0"

def create_number_input(label, key, value=0, min_value=0, max_value=100000000000, step=1000000):
    """Tạo input số với nút tăng/giảm"""
    col1, col2, col3 = st.columns([3, 1, 1])
    
    with col1:
        new_value = st.number_input(
            label,
            min_value=min_value,
            max_value=max_value,
            value=value,
            step=step,
            key=key
        )
    
    with col2:
        if st.button("➕", key=f"inc_{key}"):
            new_value += step
            st.session_state[key] = new_value
            st.rerun()
    
    with col3:
        if st.button("➖", key=f"dec_{key}"):
            new_value = max(min_value, new_value - step)
            st.session_state[key] = new_value
            st.rerun()
    
    return new_value

def display_financial_metrics(metrics):
    """Hiển thị các chỉ số tài chính"""
    if not metrics:
        st.warning("Chưa có dữ liệu để tính toán chỉ số tài chính")
        return
    
    cols = st.columns(4)
    
    with cols[0]:
        st.metric(
            "Nghĩa vụ trả nợ hàng tháng",
            f"{format_currency(metrics.get('monthly_payment', 0))} VNĐ"
        )
    
    with cols[1]:
        st.metric(
            "Tỷ lệ trả nợ (DSR)",
            f"{metrics.get('dsr_ratio', 0):.1f}%"
        )
    
    with cols[2]:
        st.metric(
            "LTV",
            f"{metrics.get('ltv', 0):.1f}%"
        )
    
    with cols[3]:
        st.metric(
            "Biên an toàn trả nợ",
            f"{metrics.get('safety_margin', 0):.1f}%"
        )

def create_payment_schedule_chart(payment_schedule):
    """Tạo biểu đồ lịch trả nợ"""
    if not payment_schedule:
        st.warning("Không có dữ liệu lịch trả nợ")
        return
    
    df = pd.DataFrame(payment_schedule)
    
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
    
    # Biểu đồ dòng tiền
    ax1.plot(df['thang'], df['goc_con_lai'], marker='o', linewidth=2)
    ax1.set_title('Dư nợ gốc theo thời gian')
    ax1.set_xlabel('Tháng')
    ax1.set_ylabel('Dư nợ gốc (VNĐ)')
    ax1.grid(True, alpha=0.3)
    ax1.tick_params(axis='x', rotation=45)
    
    # Biểu đồ phân bổ trả nợ
    months = df['thang'][::max(1, len(df)//10)]
    principal = df['tra_goc'][::max(1, len(df)//10)]
    interest = df['tra_lai'][::max(1, len(df)//10)]
    
    x = range(len(months))
    width = 0.35
    
    ax2.bar(x, principal, width, label='Trả gốc', alpha=0.7)
    ax2.bar([i + width for i in x], interest, width, label='Trả lãi', alpha=0.7)
    ax2.set_title('Phân bổ trả nợ theo tháng')
    ax2.set_xlabel('Tháng')
    ax2.set_ylabel('Số tiền (VNĐ)')
    ax2.legend()
    ax2.grid(True, alpha=0.3)
    ax2.set_xticks([i + width/2 for i in x])
    ax2.set_xticklabels(months, rotation=45)
    
    plt.tight_layout()
    st.pyplot(fig)

def create_financial_pie_chart(financial_data):
    """Tạo biểu đồ tròn phân bổ tài chính"""
    if not financial_data:
        st.warning("Không có dữ liệu để tạo biểu đồ")
        return
    
    labels = ['Vốn vay', 'Vốn đối ứng']
    sizes = [
        financial_data.get('so_tien_vay', 0),
        financial_data.get('von_doi_ung', 0)
    ]
    
    if sum(sizes) == 0:
        st.warning("Không có dữ liệu vốn")
        return
    
    colors = ['#ff9999', '#66b3ff']
    
    fig, ax = plt.subplots(figsize=(8, 8))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, colors=colors, autopct='%1.1f%%',
        startangle=90
    )
    
    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontweight('bold')
    
    ax.set_title('Phân bổ nguồn vốn')
    st.pyplot(fig)

# =============================================================================
# DOCUMENT PARSER
# =============================================================================

class DocumentParser:
    def __init__(self):
        self.patterns = {
            'ho_ten': r'Họ và tên:\s*([^\n]+)',
            'cccd': r'CMND/CCCD/hộ chiếu:\s*([^\n]+)',
            'dia_chi': r'Nơi cư trú:\s*([^\n]+)',
            'dien_thoai': r'Số điện thoại:\s*([^\n]+)',
            'tong_nhu_cau_von': r'Tổng nhu cầu vốn:\s*([\d.,]+)',
            'von_doi_ung': r'Vốn đối ứng tham gia.*?:\s*([\d.,]+)',
            'so_tien_vay': r'Vốn vay Agribank số tiền:\s*([\d.,]+)',
            'muc_dich_vay': r'Mục đích vay:\s*([^\n]+)',
            'thoi_gian_vay': r'Thời hạn vay:\s*(\d+)',
            'lai_suat': r'Lãi suất:\s*([\d.,]+)%',
            'gia_tri_tai_san': r'Giá trị.*?:\s*([\d.,]+)'
        }
    
    def parse_document(self, file):
        """Phân tích file docx và trích xuất thông tin"""
        try:
            doc = Document(file)
            full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            extracted_data = self._extract_data(full_text)
            return extracted_data
            
        except Exception as e:
            st.error(f"Lỗi khi phân tích document: {e}")
            return None
    
    def _extract_data(self, text):
        """Trích xuất dữ liệu từ text sử dụng regex patterns"""
        data = {}
        
        # Thông tin khách hàng
        customers = self._extract_customers(text)
        if customers:
            data['khach_hang'] = customers
            main_customer = customers[0]
            data.update({
                'ho_ten': main_customer['ho_ten'],
                'cccd': main_customer['cccd'],
                'dia_chi': main_customer['dia_chi'],
                'dien_thoai': main_customer['dien_thoai']
            })
        
        # Thông tin tài chính
        financial_data = self._extract_financial_info(text)
        data.update(financial_data)
        
        # Thông tin tài sản
        collateral_data = self._extract_collateral_info(text)
        data.update(collateral_data)
        
        return data
    
    def _extract_customers(self, text):
        """Trích xuất thông tin nhiều khách hàng"""
        customers = []
        
        customer_blocks = re.split(r'\d+\. Họ và tên:', text)
        
        for block in customer_blocks[1:]:
            customer = {}
            
            name_match = re.search(r'^([^-]+)', block)
            if name_match:
                customer['ho_ten'] = name_match.group(1).strip()
            
            cccd_match = re.search(self.patterns['cccd'], block)
            if cccd_match:
                customer['cccd'] = cccd_match.group(1).strip()
            
            address_match = re.search(self.patterns['dia_chi'], block)
            if address_match:
                customer['dia_chi'] = address_match.group(1).strip()
            
            phone_match = re.search(self.patterns['dien_thoai'], block)
            if phone_match:
                customer['dien_thoai'] = phone_match.group(1).strip()
            
            if customer:
                customers.append(customer)
        
        return customers
    
    def _extract_financial_info(self, text):
        """Trích xuất thông tin tài chính"""
        financial_data = {}
        
        total_match = re.search(self.patterns['tong_nhu_cau_von'], text)
        if total_match:
            financial_data['tong_nhu_cau_von'] = self._convert_currency_to_number(total_match.group(1))
        
        owner_match = re.search(self.patterns['von_doi_ung'], text, re.DOTALL)
        if owner_match:
            financial_data['von_doi_ung'] = self._convert_currency_to_number(owner_match.group(1))
        
        loan_match = re.search(self.patterns['so_tien_vay'], text)
        if loan_match:
            financial_data['so_tien_vay'] = self._convert_currency_to_number(loan_match.group(1))
        
        purpose_match = re.search(self.patterns['muc_dich_vay'], text)
        if purpose_match:
            financial_data['muc_dich_vay'] = purpose_match.group(1).strip()
        
        term_match = re.search(self.patterns['thoi_gian_vay'], text)
        if term_match:
            financial_data['thoi_gian_vay'] = int(term_match.group(1))
        
        interest_match = re.search(self.patterns['lai_suat'], text)
        if interest_match:
            financial_data['lai_suat'] = float(interest_match.group(1).replace(',', '.'))
        
        if financial_data.get('tong_nhu_cau_von') and financial_data.get('von_doi_ung'):
            financial_data['ty_le_von_doi_ung'] = (
                financial_data['von_doi_ung'] / financial_data['tong_nhu_cau_von'] * 100
            )
        
        return financial_data
    
    def _extract_collateral_info(self, text):
        """Trích xuất thông tin tài sản bảo đảm"""
        collateral_data = {}
        
        asset_match = re.search(r'Tài sản \d+.*?Giá trị.*?:\s*([\d.,]+)', text, re.DOTALL)
        if asset_match:
            collateral_data['gia_tri_thi_truong'] = self._convert_currency_to_number(asset_match.group(1))
            collateral_data['loai_tai_san'] = "Bất động sản"
        
        address_match = re.search(r'Địa chỉ.*?:\s*([^\n]+)', text)
        if address_match:
            collateral_data['dia_chi_tai_san'] = address_match.group(1).strip()
        
        if collateral_data.get('gia_tri_thi_truong') and 'so_tien_vay' in self._extract_financial_info(text):
            financial_data = self._extract_financial_info(text)
            loan_amount = financial_data.get('so_tien_vay', 0)
            asset_value = collateral_data['gia_tri_thi_truong']
            if asset_value > 0:
                collateral_data['ltv'] = (loan_amount / asset_value) * 100
        
        return collateral_data
    
    def _convert_currency_to_number(self, currency_str):
        """Chuyển đổi chuỗi tiền tệ sang số"""
        if not currency_str:
            return 0
        
        cleaned = currency_str.replace('.', '').replace(',', '.').split(' ')[0]
        
        try:
            return float(cleaned)
        except ValueError:
            return 0

# =============================================================================
# FINANCIAL CALCULATOR
# =============================================================================

class FinancialCalculator:
    def __init__(self):
        pass
    
    def calculate_payment_schedule(self, financial_data):
        """Tính toán lịch trả nợ"""
        loan_amount = financial_data.get('so_tien_vay', 0)
        interest_rate = financial_data.get('lai_suat', 0) / 100 / 12
        loan_term = financial_data.get('thoi_gian_vay', 0)
        
        if not all([loan_amount, interest_rate, loan_term]):
            return []
        
        monthly_payment = self._calculate_monthly_payment(loan_amount, interest_rate, loan_term)
        
        schedule = []
        remaining_balance = loan_amount
        
        for month in range(1, loan_term + 1):
            interest_payment = remaining_balance * interest_rate
            principal_payment = monthly_payment - interest_payment
            remaining_balance -= principal_payment
            
            if month == loan_term:
                principal_payment += remaining_balance
                remaining_balance = 0
            
            schedule.append({
                'thang': month,
                'tra_goc': round(principal_payment),
                'tra_lai': round(interest_payment),
                'tong_tra': round(principal_payment + interest_payment),
                'goc_con_lai': max(0, round(remaining_balance))
            })
        
        return schedule
    
    def _calculate_monthly_payment(self, loan_amount, monthly_rate, loan_term):
        """Tính toán khoản trả hàng tháng"""
        if monthly_rate == 0:
            return loan_amount / loan_term
        
        return loan_amount * monthly_rate * (1 + monthly_rate) ** loan_term / ((1 + monthly_rate) ** loan_term - 1)
    
    def calculate_financial_metrics(self, financial_data, customer_data):
        """Tính toán các chỉ số tài chính"""
        loan_amount = financial_data.get('so_tien_vay', 0)
        interest_rate = financial_data.get('lai_suat', 0)
        loan_term = financial_data.get('thoi_gian_vay', 0)
        asset_value = financial_data.get('gia_tri_tai_san', 0)
        
        metrics = {}
        
        if all([loan_amount, interest_rate, loan_term]):
            monthly_rate = interest_rate / 100 / 12
            monthly_payment = self._calculate_monthly_payment(loan_amount, monthly_rate, loan_term)
            metrics['monthly_payment'] = monthly_payment
        
        if asset_value > 0:
            metrics['ltv'] = (loan_amount / asset_value) * 100
        
        monthly_income = 100000000
        if 'monthly_payment' in metrics and monthly_income > 0:
            metrics['dsr_ratio'] = (metrics['monthly_payment'] / monthly_income) * 100
        
        monthly_expenses = 45000000
        if monthly_income and monthly_expenses:
            disposable_income = monthly_income - monthly_expenses
            if 'monthly_payment' in metrics and disposable_income > 0:
                metrics['safety_margin'] = ((disposable_income - metrics['monthly_payment']) / disposable_income) * 100
        
        return metrics

# =============================================================================
# DATA MANAGER
# =============================================================================

class DataManager:
    def __init__(self):
        self.customer_data = {}
        self.financial_data = {}
        self.collateral_data = {}
        self.original_data = {}
    
    def update_from_document(self, extracted_data):
        self.original_data = extracted_data.copy()
        
        if 'ho_ten' in extracted_data:
            self.customer_data = {
                'ho_ten': extracted_data.get('ho_ten', ''),
                'cccd': extracted_data.get('cccd', ''),
                'dia_chi': extracted_data.get('dia_chi', ''),
                'dien_thoai': extracted_data.get('dien_thoai', '')
            }
        
        financial_fields = [
            'tong_nhu_cau_von', 'von_doi_ung', 'so_tien_vay', 
            'ty_le_von_doi_ung', 'lai_suat', 'thoi_gian_vay', 'muc_dich_vay'
        ]
        self.financial_data = {
            field: extracted_data.get(field, 0 if field != 'muc_dich_vay' else '')
            for field in financial_fields
        }
        
        collateral_fields = [
            'loai_tai_san', 'gia_tri_thi_truong', 'dia_chi_tai_san', 'ltv', 'giay_to_phap_ly'
        ]
        self.collateral_data = {
            field: extracted_data.get(field, 0 if field not in ['loai_tai_san', 'dia_chi_tai_san', 'giay_to_phap_ly'] else '')
            for field in collateral_fields
        }
    
    def update_customer_data(self, data):
        self.customer_data.update(data)
    
    def update_financial_data(self, data):
        self.financial_data.update(data)
    
    def update_collateral_data(self, data):
        self.collateral_data.update(data)
    
    def get_customer_data(self):
        return self.customer_data.copy()
    
    def get_financial_data(self):
        return self.financial_data.copy()
    
    def get_collateral_data(self):
        return self.collateral_data.copy()
    
    def get_original_data(self):
        return self.original_data.copy()

# =============================================================================
# GEMINI CLIENT (SIMPLIFIED)
# =============================================================================

class GeminiClient:
    def __init__(self):
        self.api_key = None
        self.configured = False
    
    def set_api_key(self, api_key):
        self.api_key = api_key
        if api_key:
            self.configured = True
    
    def is_configured(self):
        return self.configured
    
    def analyze_financial_data(self, data, data_source):
        if not self.is_configured():
            return "Vui lòng nhập API key Google AI Studio ở sidebar"
        
        analysis = f"""
PHÂN TÍCH TÀI CHÍNH - NGUỒN DỮ LIỆU: {data_source}

ĐÁNH GIÁ RỦI RO:
• Khả năng trả nợ: {data.get('metrics', {}).get('dsr_ratio', 0):.1f}% - {'Tốt' if data.get('metrics', {}).get('dsr_ratio', 0) < 40 else 'Cần thận trọng'}
• Tỷ lệ LTV: {data.get('metrics', {}).get('ltv', 0):.1f}% - {'An toàn' if data.get('metrics', {}).get('ltv', 0) < 80 else 'Cao'}
• Biên an toàn: {data.get('metrics', {}).get('safety_margin', 0):.1f}%

ĐỀ XUẤT:
• Xem xét khả năng trả nợ dựa trên thu nhập ổn định
• Đánh giá tính khả thi của phương án sử dụng vốn
• Kiểm tra tính pháp lý của tài sản bảo đảm

LƯU Ý: Đây là phân tích mẫu. Vui lòng tích hợp API key thực tế để có phân tích chi tiết từ Gemini AI.
"""
        return analysis
    
    def chat(self, message):
        if not self.is_configured():
            return "Vui lòng nhập API key Google AI Studio ở sidebar để sử dụng tính năng chat"
        
        return f"Tính năng chat với Gemini AI sẽ hoạt động khi bạn tích hợp API key thực tế. Câu hỏi của bạn: '{message}'"

# =============================================================================
# EXPORTERS
# =============================================================================

class ExcelExporter:
    def export_payment_schedule(self, payment_schedule):
        df = pd.DataFrame(payment_schedule)
        
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='KeHoachTraNo', index=False)
        
        output.seek(0)
        return output.getvalue()

class ReportExporter:
    def export_word_report(self, data, include_charts=True):
        doc = Document()
        
        title = doc.add_heading('BÁO CÁO THẨM ĐỊNH TÍN DỤNG', 0)
        title.alignment = 1
        
        doc.add_heading('I. THÔNG TIN KHÁCH HÀNG', level=1)
        customer = data.get('customer', {})
        doc.add_paragraph(f"Họ và tên: {customer.get('ho_ten', '')}")
        doc.add_paragraph(f"CCCD/CMND: {customer.get('cccd', '')}")
        doc.add_paragraph(f"Địa chỉ: {customer.get('dia_chi', '')}")
        doc.add_paragraph(f"Số điện thoại: {customer.get('dien_thoai', '')}")
        
        doc.add_heading('II. THÔNG TIN TÀI CHÍNH', level=1)
        financial = data.get('financial', {})
        doc.add_paragraph(f"Tổng nhu cầu vốn: {format_currency(financial.get('tong_nhu_cau_von', 0))} VNĐ")
        doc.add_paragraph(f"Vốn đối ứng: {format_currency(financial.get('von_doi_ung', 0))} VNĐ")
        doc.add_paragraph(f"Số tiền vay: {format_currency(financial.get('so_tien_vay', 0))} VNĐ")
        doc.add_paragraph(f"Lãi suất: {financial.get('lai_suat', 0)}%/năm")
        doc.add_paragraph(f"Thời gian vay: {financial.get('thoi_gian_vay', 0)} tháng")
        
        doc.add_heading('III. CHỈ SỐ TÀI CHÍNH', level=1)
        metrics = data.get('metrics', {})
        doc.add_paragraph(f"Nghĩa vụ trả nợ hàng tháng: {format_currency(metrics.get('monthly_payment', 0))} VNĐ")
        doc.add_paragraph(f"Tỷ lệ trả nợ (DSR): {metrics.get('dsr_ratio', 0):.1f}%")
        doc.add_paragraph(f"LTV: {metrics.get('ltv', 0):.1f}%")
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    def export_pdf_report(self, data, include_charts=True):
        # Simplified PDF export - returns Word document as fallback
        return self.export_word_report(data, include_charts)

# =============================================================================
# TAB MANAGEMENT
# =============================================================================

def create_sidebar():
    with st.sidebar:
        st.header("🔑 Cài đặt API")
        
        api_key = st.text_input(
            "Google AI Studio API Key",
            type="password",
            help="Nhập API key từ Google AI Studio để sử dụng Gemini AI"
        )
        
        if api_key:
            st.session_state.gemini_client.set_api_key(api_key)
            st.success("✅ API key đã được thiết lập")
        
        st.markdown("---")
        st.header("📤 Upload File")
        
        uploaded_file = st.file_uploader(
            "Tải lên file PASDV.docx",
            type=['docx'],
            help="Tải lên file phương án sử dụng vốn định dạng .docx"
        )
        
        if uploaded_file is not None:
            try:
                parser = DocumentParser()
                extracted_data = parser.parse_document(uploaded_file)
                
                if extracted_data:
                    st.session_state.data_manager.update_from_document(extracted_data)
                    st.success("✅ File đã được xử lý thành công!")
                    
                    with st.expander("📋 Xem thông tin trích xuất từ file"):
                        if 'khach_hang' in extracted_data:
                            for kh in extracted_data['khach_hang']:
                                st.write(f"**{kh['ho_ten']}** - {kh['cccd']}")
                else:
                    st.error("❌ Không thể trích xuất dữ liệu từ file")
            except Exception as e:
                st.error(f"❌ Lỗi khi xử lý file: {str(e)}")
        
        st.markdown("---")
        st.header("💡 Hướng dẫn")
        st.info("""
        1. Nhập API key Google AI Studio
        2. Upload file PASDV.docx
        3. Kiểm tra và chỉnh sửa dữ liệu ở các tab
        4. Phân tích với AI và xuất báo cáo
        """)

def create_customer_info_tab():
    st.header("👤 Thông Tin Định Danh Khách Hàng")
    
    data_manager = st.session_state.data_manager
    customer_data = data_manager.get_customer_data()
    
    col1, col2 = st.columns(2)
    
    with col1:
        ho_ten = st.text_input("Họ và tên", value=customer_data.get('ho_ten', ''), key="customer_name")
        cccd = st.text_input("CCCD/CMND", value=customer_data.get('cccd', ''), key="customer_id")
    
    with col2:
        dia_chi = st.text_input("Địa chỉ", value=customer_data.get('dia_chi', ''), key="customer_address")
        dien_thoai = st.text_input("Số điện thoại", value=customer_data.get('dien_thoai', ''), key="customer_phone")
    
    if st.button("💾 Lưu thông tin khách hàng"):
        updated_data = {
            'ho_ten': ho_ten,
            'cccd': cccd,
            'dia_chi': dia_chi,
            'dien_thoai': dien_thoai
        }
        data_manager.update_customer_data(updated_data)
        st.success("✅ Thông tin khách hàng đã được cập nhật")

def create_financial_info_tab():
    st.header("💰 Thông Tin Tài Chính / Phương Án Sử Dụng Vốn")
    
    data_manager = st.session_state.data_manager
    financial_data = data_manager.get_financial_data()
    
    col1, col2 = st.columns(2)
    
    with col1:
        muc_dich_vay = st.text_area("Mục đích vay", value=financial_data.get('muc_dich_vay', ''), height=100, key="loan_purpose")
        tong_nhu_cau_von = create_number_input("Tổng nhu cầu vốn (VNĐ)", "total_capital_needed", value=financial_data.get('tong_nhu_cau_von', 0))
        von_doi_ung = create_number_input("Vốn đối ứng (VNĐ)", "owner_capital", value=financial_data.get('von_doi_ung', 0))
    
    with col2:
        so_tien_vay = create_number_input("Số tiền vay (VNĐ)", "loan_amount", value=financial_data.get('so_tien_vay', 0))
        lai_suat = st.number_input("Lãi suất vay (%/năm)", min_value=0.0, max_value=50.0, value=financial_data.get('lai_suat', 0.0), step=0.1, key="interest_rate")
        thoi_gian_vay = st.number_input("Thời gian vay (tháng)", min_value=1, max_value=360, value=financial_data.get('thoi_gian_vay', 0), key="loan_term")
    
    if st.button("💾 Lưu thông tin tài chính"):
        updated_data = {
            'muc_dich_vay': muc_dich_vay,
            'tong_nhu_cau_von': tong_nhu_cau_von,
            'von_doi_ung': von_doi_ung,
            'so_tien_vay': so_tien_vay,
            'lai_suat': lai_suat,
            'thoi_gian_vay': thoi_gian_vay
        }
        data_manager.update_financial_data(updated_data)
        st.success("✅ Thông tin tài chính đã được cập nhật")

def create_collateral_tab():
    st.header("🏠 Tài Sản Bảo Đảm")
    
    data_manager = st.session_state.data_manager
    collateral_data = data_manager.get_collateral_data()
    
    col1, col2 = st.columns(2)
    
    with col1:
        loai_tai_san = st.selectbox("Loại tài sản", ["Bất động sản", "Xe ô tô", "Thiết bị máy móc", "Tài sản khác"], key="asset_type")
        gia_tri_thi_truong = create_number_input("Giá trị thị trường (VNĐ)", "market_value", value=collateral_data.get('gia_tri_thi_truong', 0))
    
    with col2:
        dia_chi_tai_san = st.text_input("Địa chỉ tài sản", value=collateral_data.get('dia_chi_tai_san', ''), key="asset_address")
        giay_to_phap_ly = st.text_area("Giấy tờ pháp lý", value=collateral_data.get('giay_to_phap_ly', ''), height=100, key="legal_docs")
    
    if st.button("💾 Lưu thông tin tài sản"):
        updated_data = {
            'loai_tai_san': loai_tai_san,
            'gia_tri_thi_truong': gia_tri_thi_truong,
            'dia_chi_tai_san': dia_chi_tai_san,
            'giay_to_phap_ly': giay_to_phap_ly
        }
        data_manager.update_collateral_data(updated_data)
        st.success("✅ Thông tin tài sản đã được cập nhật")

def create_financial_calculation_tab():
    st.header("📊 Tính Toán Chỉ Tiêu Tài Chính / Dòng Tiền")
    
    data_manager = st.session_state.data_manager
    financial_data = data_manager.get_financial_data()
    customer_data = data_manager.get_customer_data()
    
    if not financial_data.get('so_tien_vay') or not financial_data.get('lai_suat'):
        st.warning("Vui lòng nhập đầy đủ thông tin tài chính ở tab trước")
        return
    
    calculator = FinancialCalculator()
    metrics = calculator.calculate_financial_metrics(financial_data, customer_data)
    payment_schedule = calculator.calculate_payment_schedule(financial_data)
    
    display_financial_metrics(metrics)
    
    st.subheader("📋 Kế hoạch trả nợ")
    
    if payment_schedule:
        df = pd.DataFrame(payment_schedule)
        st.dataframe(df, use_container_width=True)
        
        st.session_state.payment_schedule = payment_schedule
        st.session_state.financial_metrics = metrics

def create_charts_tab():
    st.header("📈 Biểu Đồ Phân Tích Tài Chính")
    
    data_manager = st.session_state.data_manager
    financial_data = data_manager.get_financial_data()
    
    if not financial_data:
        st.warning("Chưa có dữ liệu tài chính để vẽ biểu đồ")
        return
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Phân bổ nguồn vốn")
        create_financial_pie_chart(financial_data)
    
    with col2:
        st.subheader("Lịch trả nợ")
        payment_schedule = getattr(st.session_state, 'payment_schedule', [])
        create_payment_schedule_chart(payment_schedule)

def create_ai_analysis_tab():
    st.header("🤖 Phân Tích AI Gemini")
    
    if not st.session_state.gemini_client.is_configured():
        st.warning("Vui lòng nhập API key ở sidebar để sử dụng tính năng AI")
        return
    
    data_manager = st.session_state.data_manager
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📄 Phân tích từ File Upload")
        if st.button("🔍 Phân tích dữ liệu gốc", key="analyze_original"):
            with st.spinner("AI đang phân tích dữ liệu từ file..."):
                original_data = data_manager.get_original_data()
                analysis = st.session_state.gemini_client.analyze_financial_data(
                    {'original_data': original_data, 'source': 'file_upload'}, 
                    "dữ liệu gốc từ file upload"
                )
                st.text_area("Kết quả phân tích", analysis, height=300, key="analysis_original")
    
    with col2:
        st.subheader("✏️ Phân tích dữ liệu đã chỉnh sửa")
        if st.button("🔍 Phân tích dữ liệu hiện tại", key="analyze_current"):
            with st.spinner("AI đang phân tích dữ liệu hiện tại..."):
                current_data = {
                    'customer': data_manager.get_customer_data(),
                    'financial': data_manager.get_financial_data(),
                    'collateral': data_manager.get_collateral_data(),
                    'metrics': getattr(st.session_state, 'financial_metrics', {})
                }
                analysis = st.session_state.gemini_client.analyze_financial_data(
                    current_data, "dữ liệu sau khi hiệu chỉnh tại giao diện"
                )
                st.text_area("Kết quả phân tích", analysis, height=300, key="analysis_current")

def create_chatbox_tab():
    st.header("💬 Chatbox Gemini")
    
    if not st.session_state.gemini_client.is_configured():
        st.warning("Vui lòng nhập API key ở sidebar để sử dụng chatbox")
        return
    
    # Hiển thị lịch sử chat
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.write(message["content"])
    
    # Input chat
    prompt = st.chat_input("Nhập câu hỏi của bạn...")
    
    if prompt:
        # Thêm câu hỏi vào lịch sử
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        
        with st.chat_message("user"):
            st.write(prompt)
        
        # Nhận phản hồi từ AI
        with st.spinner("AI đang suy nghĩ..."):
            response = st.session_state.gemini_client.chat(prompt)
            st.session_state.chat_history.append({"role": "assistant", "content": response})
        
        with st.chat_message("assistant"):
            st.write(response)
    
    # Nút xóa hội thoại
    if st.button("🗑️ Xóa hội thoại"):
        st.session_state.chat_history = []
        st.rerun()

def create_export_tab():
    st.header("📤 Xuất File Báo Cáo")
    
    data_manager = st.session_state.data_manager
    
    export_option = st.selectbox(
        "Chọn loại file xuất",
        [
            "Xuất bảng kê kế hoạch trả nợ (Excel)",
            "Xuất báo cáo thẩm định (Word/PDF)"
        ]
    )
    
    if export_option == "Xuất bảng kê kế hoạch trả nợ (Excel)":
        if st.button("📊 Xuất file Excel"):
            payment_schedule = getattr(st.session_state, 'payment_schedule', [])
            if payment_schedule:
                exporter = ExcelExporter()
                excel_file = exporter.export_payment_schedule(payment_schedule)
                
                st.download_button(
                    label="📥 Tải xuống file Excel",
                    data=excel_file,
                    file_name="ke_hoach_tra_no.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            else:
                st.warning("Không có dữ liệu kế hoạch trả nợ để xuất")
    
    else:  # Xuất báo cáo thẩm định
        col1, col2 = st.columns(2)
        
        with col1:
            report_type = st.radio(
                "Định dạng báo cáo",
                ["Word (.docx)", "PDF (.pdf)"]
            )
        
        with col2:
            include_charts = st.checkbox("Bao gồm biểu đồ", value=True)
        
        if st.button("📄 Tạo báo cáo thẩm định"):
            with st.spinner("Đang tạo báo cáo..."):
                exporter = ReportExporter()
                
                # Thu thập dữ liệu cho báo cáo
                report_data = {
                    'customer': data_manager.get_customer_data(),
                    'financial': data_manager.get_financial_data(),
                    'collateral': data_manager.get_collateral_data(),
                    'metrics': getattr(st.session_state, 'financial_metrics', {}),
                    'payment_schedule': getattr(st.session_state, 'payment_schedule', [])
                }
                
                if report_type == "Word (.docx)":
                    report_file = exporter.export_word_report(report_data, include_charts)
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    file_name = "bao_cao_tham_dinh.docx"
                else:
                    report_file = exporter.export_pdf_report(report_data, include_charts)
                    mime_type = "application/pdf"
                    file_name = "bao_cao_tham_dinh.pdf"
                
                st.download_button(
                    label=f"📥 Tải xuống {file_name}",
                    data=report_file,
                    file_name=file_name,
                    mime=mime_type
                )

def create_tabs():
    """Tạo các tab chính của ứng dụng"""
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
        "👤 Thông Tin Khách Hàng",
        "💰 Thông Tin Tài Chính", 
        "🏠 Tài Sản Bảo Đảm",
        "📊 Tính Toán Tài Chính",
        "📈 Biểu Đồ",
        "🤖 Phân Tích AI",
        "💬 Chatbox Gemini",
        "📤 Xuất File"
    ])
    
    with tab1:
        create_customer_info_tab()
    
    with tab2:
        create_financial_info_tab()
    
    with tab3:
        create_collateral_tab()
    
    with tab4:
        create_financial_calculation_tab()
    
    with tab5:
        create_charts_tab()
    
    with tab6:
        create_ai_analysis_tab()
    
    with tab7:
        create_chatbox_tab()
    
    with tab8:
        create_export_tab()

# =============================================================================
# MAIN APPLICATION
# =============================================================================

def main():
    # Cấu hình trang
    st.set_page_config(
        page_title="CADAP - Credit Analysis & Decision Assistance Platform",
        page_icon="🏦",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Khởi tạo session state
    if 'data_manager' not in st.session_state:
        st.session_state.data_manager = DataManager()
    if 'gemini_client' not in st.session_state:
        st.session_state.gemini_client = GeminiClient()
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    
    # Tạo sidebar
    create_sidebar()
    
    # Header
    st.title("🏦 CADAP - Credit Analysis & Decision Assistance Platform")
    st.markdown("---")
    
    # Tạo các tab
    create_tabs()

if __name__ == "__main__":
    main()
